#!/usr/bin/env python3
"""
Examine Week 1 Day 1 DOCX structure to understand content for workflow design.
"""

from docx import Document
import json
from pathlib import Path


def analyze_document_structure(doc):
    """Analyze the overall document structure."""
    structure = {
        'total_paragraphs': len(doc.paragraphs),
        'total_tables': len(doc.tables),
        'styles_used': set(),
        'headings': []
    }

    for para in doc.paragraphs:
        style_name = para.style.name
        structure['styles_used'].add(style_name)

        if 'Heading' in style_name or style_name == 'Title':
            structure['headings'].append({
                'style': style_name,
                'text': para.text
            })

    structure['styles_used'] = sorted(list(structure['styles_used']))
    return structure


def extract_day_content(doc, day_number=1):
    """Extract content for a specific day."""
    day_content = {
        'day_number': day_number,
        'day_title': None,
        'sections': []
    }

    current_section = None
    in_target_day = False

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        # Check if we're entering the target day
        if f'Day {day_number}' in text and ('Heading' in style or style == 'Title'):
            in_target_day = True
            day_content['day_title'] = text
            continue

        # Check if we've moved to next day
        if in_target_day and any(f'Day {i}' in text for i in range(day_number + 1, 10)):
            break

        if not in_target_day:
            continue

        # Process content within target day
        if style.startswith('Heading'):
            # Start new section
            if current_section:
                day_content['sections'].append(current_section)
            current_section = {
                'heading_level': style,
                'heading_text': text,
                'paragraphs': []
            }
        elif style == 'Normal' or 'Body' in style:
            # Add paragraph to current section
            if current_section is not None:
                if text:  # Only add non-empty paragraphs
                    current_section['paragraphs'].append(text)
            elif text:
                # Paragraph before any section heading
                if not day_content['sections']:
                    day_content['sections'].append({
                        'heading_level': 'None',
                        'heading_text': 'Introduction',
                        'paragraphs': [text]
                    })

    # Add last section
    if current_section:
        day_content['sections'].append(current_section)

    return day_content


def extract_tables_content(doc):
    """Extract all tables from document."""
    tables_data = []

    for i, table in enumerate(doc.tables):
        table_data = {
            'table_index': i,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'data': []
        }

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data['data'].append(row_data)

        tables_data.append(table_data)

    return tables_data


def main():
    # Path to Week 1 document
    docx_path = Path('/app/FINAL INTEGRO CONTENT/FINAL INTEGRO CONTENT/Expedition 1_ Healing the Self/Week 1_ Roots of Healing/Week 1_ Roots of Healing.docx')

    if not docx_path.exists():
        print(f"Error: File not found at {docx_path}")
        return

    print(f"Loading document: {docx_path.name}")
    doc = Document(str(docx_path))

    # Analyze overall structure
    print("\n" + "="*80)
    print("DOCUMENT STRUCTURE ANALYSIS")
    print("="*80)
    structure = analyze_document_structure(doc)
    print(json.dumps(structure, indent=2, default=str))

    # Extract Day 1 content
    print("\n" + "="*80)
    print("DAY 1 CONTENT EXTRACTION")
    print("="*80)
    day1_content = extract_day_content(doc, day_number=1)
    print(json.dumps(day1_content, indent=2))

    # Extract tables
    print("\n" + "="*80)
    print("TABLES IN DOCUMENT")
    print("="*80)
    tables = extract_tables_content(doc)
    print(json.dumps(tables, indent=2))

    # Save to files for review
    output_dir = Path('/app/scripts/output')
    output_dir.mkdir(exist_ok=True, parents=True)

    with open(output_dir / 'week1_structure.json', 'w') as f:
        json.dump(structure, f, indent=2, default=str)

    with open(output_dir / 'week1_day1_content.json', 'w') as f:
        json.dump(day1_content, f, indent=2)

    with open(output_dir / 'week1_tables.json', 'w') as f:
        json.dump(tables, f, indent=2)

    print("\n" + "="*80)
    print("OUTPUT SAVED")
    print("="*80)
    print(f"Structure: {output_dir / 'week1_structure.json'}")
    print(f"Day 1 Content: {output_dir / 'week1_day1_content.json'}")
    print(f"Tables: {output_dir / 'week1_tables.json'}")


if __name__ == '__main__':
    main()
